import spacy
import json
import openai
import os
import requests
import time
import firebase_admin
import datetime
from google.oauth2 import service_account
from google.cloud import secretmanager, tasks_v2
from google.protobuf import timestamp_pb2
from firebase_admin import credentials, firestore
from io import BytesIO
from langdetect import detect
from flask import Flask, jsonify, request
from flask_cors import CORS
from flask_mail import Mail, Message
from dotenv import load_dotenv, find_dotenv
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH

app = Flask(__name__)
CORS(app, resources={
    r"/*": {
        "origins": ["https://myaiguy.net", "https://www.myaiguy.net", "http://localhost:3000"]
    }
})

app.config["ENV"] = "development"
load_dotenv(find_dotenv(usecwd=True))

def access_secret_version(secret_id, project_id="coverlettergenerator-396114", version_id="latest"):
    secrets_key_file = 'CLRSecretCreds.json'
    # Access the payload of the secret version.
    client = secretmanager.SecretManagerServiceClient(credentials=service_account.Credentials.from_service_account_file(secrets_key_file))

    # Build the resource name of the secret version
    name = f"projects/{project_id}/secrets/{secret_id}/versions/{version_id}"

    # Access the secret version
    response = client.access_secret_version(name=name)

    # Return the decoded payload
    return response.payload.data.decode('UTF-8')

# Key Variables
MAILGUN_DOMAIN = 'mg.myaiguy.net'
MAILGUN_API_KEY = access_secret_version('MAILGUN_API_KEY')
MAILGUN_API_ENDPOINT = f"https://api.mailgun.net/v3/mg.myaiguy.net/messages"
openai.api_key = access_secret_version('OPENAI_API_KEY')
DEEPL_API_KEY = access_secret_version('DEEPL_API_KEY')
DEEPL_API_ENDPOINT = "https://api-free.deepl.com/v2/translate"

# Get Firebase service account from Secret Manager
firebase_service_account_str = access_secret_version("Firebase-service-account")
firebase_service_account_dict = json.loads(firebase_service_account_str)

# Use the dictionary as credentials for Firebase
cred = credentials.Certificate(firebase_service_account_dict)
firebase_admin.initialize_app(cred)
db = firestore.client()

mail = Mail(app)
nlp = spacy.load("spacy_models/en_core_web_sm/en_core_web_sm-3.6.0")
driver_path = "/chromedriver.exe"

def score_content(text):
    keywords = ["about", "mission", "vision", "history", "company", "founded", "established", "believe", "our goal",
                "we aim", "we strive"]
    doc = nlp(text)

    # Score based on keywords and named entities
    keyword_score = sum([text.lower().count(keyword) for keyword in keywords])
    named_entity_score = len([ent for ent in doc.ents if ent.label_ in ["ORG", "DATE"]])

    return keyword_score + named_entity_score

def extract_relevant_content(text_content):
    # Split the text into sentences by breaking at each period
    sentences = [sent.strip() for sent in text_content.split('.') if sent]

    # Score each sentence
    scored_sentences = [(sent, score_content(sent)) for sent in sentences]

    # Sort sentences based on score and length
    scored_sentences.sort(key=lambda s: (-s[1], len(s[0])))

    # Extract the top sentences (adjust this based on your requirements)
    top_sentences = [s[0] for s in scored_sentences[:10]]

    return ' '.join(top_sentences)

def get_web_content(url):
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3"
        }
        response = requests.get(url, headers=headers)
        print("Response code for website URL:", response.status_code)
        response.raise_for_status()
        html_content = response.text
        soup = BeautifulSoup(html_content, 'html.parser')
        paragraphs = soup.find_all('p')
        content = " ".join([para.text for para in paragraphs[:10]])
        if len(soup.get_text()) > 50:
            return content
    except requests.exceptions.RequestException as e:
        print(f"Error fetching the page with requests: {e}")

    # If the requests approach fails, fall back to Selenium
    options = Options()
    options.add_argument("--headless")
    chrome_prefs = {
        "profile.default_content_settings": {"images": 2},
        "profile.managed_default_content_settings": {"images": 2}
    }
    options.experimental_options["prefs"] = chrome_prefs
    browser = webdriver.Chrome(options=options)

    try:
        print("Fetching content from:", url)
        browser.get(url)
        WebDriverWait(browser, 10).until(EC.presence_of_element_located((By.TAG_NAME, 'p')))
        paragraphs = browser.find_elements(By.TAG_NAME, 'p')
        content = " ".join([para.text for para in paragraphs[:10]])
        print("Fetched content")
    finally:
        browser.quit()

    return content

# def extract_relevant_content(html_content):
#     soup = BeautifulSoup(html_content, 'html.parser')
#
#     # Remove these tags to clean the content
#     for script in soup(["script", "style", "head", "footer", "nav"]):
#         script.extract()
#
#     # Get all the text from the cleaned HTML
#     full_text = soup.get_text()
#
#     print("Full Text after cleaning:")
#     print(full_text[:1000])
#
#     # Split the full text into sentences by breaking at each period
#     sentences = [sent.strip() for sent in full_text.split('.') if sent]
#     print("\nSentences extracted:")
#     for s in sentences[:10]:  # Printing the first 10 sentences for brevity. Adjust as needed.
#         print(s)
#
#     # Score each sentence
#     scored_sentences = [(sent, score_content(sent)) for sent in sentences]
#
#     # Sort sentences based on score and length
#     scored_sentences.sort(key=lambda s: (-s[1], len(s[0])))
#
#     # Extract the top sentences (adjust this based on your requirements)
#     top_sentences = [s[0] for s in scored_sentences[:10]]
#
#     print("\nTop Sentences after scoring:")
#     for sent in top_sentences:
#         print(sent)
#
#     return ' '.join(top_sentences)

# def get_web_content(url):
#     # Set up the Selenium Chrome driver
#     # Make sure the chromedriver executable is in the same directory as your script or provide the full path
#     options = webdriver.ChromeOptions()
#
#     # Prevent loading images
#     chrome_prefs = {}
#     options.experimental_options["prefs"] = chrome_prefs
#     chrome_prefs["profile.default_content_settings"] = {"images": 2}
#     chrome_prefs["profile.managed_default_content_settings"] = {"images": 2}
#
#     options.add_argument("--headless=new")
#     browser = webdriver.Chrome(options=options)
#
#     try:
#         print("Fetching content from:", url)
#         browser.get(url)
#
#         # Wait for at least one paragraph to be present
#         WebDriverWait(browser, 10).until(EC.presence_of_element_located((By.TAG_NAME, 'p')))
#
#         # Get all the text inside paragraph tags
#         paragraphs = browser.find_elements(By.TAG_NAME, 'p')
#         content = " ".join([para.text for para in paragraphs[:10]])
#
#         print("Fetched content")
#     finally:
#         # It's good practice to close the browser when done
#         browser.quit()
#
#     return content

def check_missing_fields(data, required_fields):
    missing = [field for field in required_fields if not data.get(field)]
    return missing

def clean_response_text(response_text):
    # List of common salutations that might be used.
    salutations = [
        "Dear Hiring Manager,",
        "To whom it may concern,",
        "Hello,",
        "Hi,",
        "Greetings,"
        "Hiring Manager,",
        "Dear [Hiring Manager&#39;s Name],",
    ]

    # List of common valedictions (closings) that might be used.
    closings = [
        "Sincerely",
        "Best regards",
        "Warm regards",
        "Kind regards",
        "Thank you",
        "Cheers",
        "Yours truly",
        "Yours sincerely",
        "Yours faithfully",
        "Regards",
        "Best",
        "In conclusion",
        "In summary",
    ]

    # Remove any salutations from the start of the text.
    for salutation in salutations:
        response_text = response_text.replace(salutation, '\n\n').strip()

    # Heuristic to remove signatures and closings.
    for closing in closings:
        if closing in response_text:
            response_text = response_text.split(closing)[0].strip() + '\n'

    return response_text

def safe_openai_request(request_function, max_retries=3, delay=5):
    """
    Safe wrapper around OpenAI requests with retries.

    Args:
    - request_function: A function that makes the actual OpenAI request.
    - max_retries: Maximum number of times to retry the request.
    - delay: How many seconds to wait between retries.

    Returns:
    - The response from the OpenAI request.
    """
    for attempt in range(max_retries):
        try:
            return request_function()
        except openai.error.ServiceUnavailableError:
            if attempt < max_retries - 1:  # i.e. not the last attempt
                time.sleep(delay)
                continue
            else:
                raise


def translate_text(text, target_language="EN"):
    try:
        detected_language = detect(text)
        # If detected language is English, return the text without translating
        if detected_language == 'en':
            return text
    except:
        # If there's any issue detecting the language, proceed with translation
        pass

    headers = {
        "Authorization": f"DeepL-Auth-Key {DEEPL_API_KEY}",
        "Content-Type": "application/json",
        "User-Agent": "CoverLetterRewriterApp/1.0"
    }

    payload = {
        "text": [text],
        "target_lang": target_language,
    }

    response = requests.post(DEEPL_API_ENDPOINT, headers=headers, json=payload)

    # Print debugging information
    print("DeepL API Status Code:", response.status_code)

    try:
        response_data = response.json()
        return response_data['translations'][0]['text']
    except Exception as e:
        print(f"Failed to parse JSON: {e}")
        return None

def rewrite_intro(name, role, company, intro_points):
    translated_intro_points = translate_text(intro_points)

    # First OpenAI call
    response = safe_openai_request(lambda: openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system",
             "content": "You are a wise and experienced career coach with a specialty in resumes and cover letters"},
            {"role": "user",
             "content": f"We must write an intro paragraph for a cover letter. Please write ONLY ONE PARAGRAPH. I will give you a few key points about myself and you must incorporate them and also highlight why I would be a strong fit for the position. This is for a very important job. This is only the intro, so please DO NOT ADD SIGNATURE AT THE END. The intro paragraph should stand out among a crowded field, and it should generate curiosity and intrigue to the first human gatekeeper. Please write in a human style. My name is {name}, and I am applying for {role} at {company}. Here are the bullet points: {translated_intro_points}"},
        ]
    ))
    initial_response_text = response['choices'][0]['message']['content']

    # Extract the main content and send it back for revision
    main_content = initial_response_text.strip()
    # Remove the "Dear Hiring Manager," line
    if ',' in main_content:
        main_content = main_content.split(',', 1)[1].strip()
    # Second OpenAI call
    revised_response = safe_openai_request(lambda: openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system",
             "content": "You are a meticulous editor. Please revise the following content for clarity and impact."},
            {"role": "user",
             "content": f"I am going to give you an intro paragraph for a cover letter. It's pretty good, but seems robotic, and we need to humanize it while remaining professional. We are looking for a very assertive tone. Please ONLY RETURN TWO PARAGRAPHS AND NO MORE. This is only the intro, so please DO NOT ADD SIGNATURE AT THE END. My name is {name} and I am applying for a {role} position for {company}. Here is the paragraph {main_content}"},
        ]
    ))
    cleaned_response = clean_response_text(revised_response['choices'][0]['message']['content'].strip())
    print(cleaned_response)
    return cleaned_response

def rewrite_tailored_experience(role, company, tailored_experience_points):
    translated_tailored_experience_points = translate_text(tailored_experience_points)

    # First OpenAI call
    response = safe_openai_request(lambda: openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system",
             "content": "You are a wise and experienced career coach with a specialty in resumes and cover letters"},
            {"role": "user",
             "content": f"We need to write a tailored experience section for a cover letter. Please write ONLY TWO PARAGRAPHS. I will give you some key points about my experience and skills and you must incorporate them to highlight why I'd be a great fit for the role. This is for a significant role and it's crucial that this section showcases my expertise effectively. Please write in a compelling and assertive style. Again, I am applying for {role} at {company}. Here are the bullet points: {translated_tailored_experience_points}"},
        ]
    ))
    initial_response_text = response['choices'][0]['message']['content']

    # Extract the main content and send it back for revision
    main_content = initial_response_text.strip()

    # Second OpenAI call for revision
    revised_response = safe_openai_request(lambda: openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system",
             "content": "You are a meticulous editor. Please revise the following content for clarity, conciseness, and impact."},
            {"role": "user",
             "content": f"I'm sharing a tailored experience section of a cover letter with you. Please write ONLY TWO PARAGRAPHS. I think it's decent, but I'm confident it can be better. Please rewrite it with a more humanized touch, ensuring it stands out from other applications and effectively showcases my skills and experience. Here is the paragraph: {main_content}"},
        ]
    ))
    cleaned_revised_response = clean_response_text(revised_response['choices'][0]['message']['content'].strip())
    print(cleaned_revised_response)
    return cleaned_revised_response

def rewrite_company_alignment(company_url, name, company):
    text_content = None
    if company_url:  # Checking if the URL is provided
        try:
            # Step 1: Fetch the HTML content from the company URL
            text_content = get_web_content(company_url)
        except Exception as e:
            print(f"Failed to fetch content from URL due to: {e}")

    if not text_content:
        # Generic OpenAI call for a broad alignment with companies
        response = safe_openai_request(lambda: openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": "You are a talented writer, experienced in aligning professional strengths with company values in cover letters"},
                {"role": "user",
                 "content": f"I don't have specific information about the company {company}. However, this is the final section of the letter, so DO NOT USE SALUTATIONS, but include a signature that ends with {name}. Can you craft ONE OR TWO paragraphs for my cover letter that emphasizes a generic alignment with popular company values like innovation, dedication, teamwork, and excellence?"}
            ]
        ))
    else:
        # Step 2: Extract the most relevant content using the score_content function
        relevant_content = extract_relevant_content(text_content)

        # Step 3: Send the extracted content to OpenAI for rephrasing
        response = safe_openai_request(lambda: openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": "You are a talented writer, experienced in aligning professional strengths with company values in cover letters"},
                {"role": "user",
                 "content": f"I've extracted some content about the company {company} from their website. This is the final section of the letter, so DO NOT USE SALUTATIONS, but include a signature that ends with {name}. Can you help me craft ONE OR TWO paragraphs for my cover letter that emphasizes my alignment with the company's values and mission derived from the main points that I will provide? Here's the extracted content: {relevant_content}"}
            ]
        ))

    alignment_paragraph_raw = response['choices'][0]['message']['content'].strip()
    alignment_paragraph_cleaned = clean_response_text(alignment_paragraph_raw)
    print("Company Alignment Content: ", alignment_paragraph_cleaned)
    return alignment_paragraph_cleaned

@app.route('/test', methods=['GET'])
def test():
    return jsonify({"message": "Server is running!"})


@app.route('/enqueue-cover-letter-task', methods=['POST'])
def enqueue_task():
    tasks_key_file = 'cloudrun-deploy-account.json'
    client = tasks_v2.CloudTasksClient(credentials=service_account.Credentials.from_service_account_file(tasks_key_file))

    PROJECT_ID = 'coverlettergenerator-396114'
    LOCATION = 'us-east1'
    QUEUE_NAME = 'cover-letter-queue'

    # Define the queue name and location
    parent = client.queue_path(PROJECT_ID, LOCATION, QUEUE_NAME)

    # Create a task
    task = {
        'app_engine_http_request': {
            'http_method': 'POST',
            'relative_uri': '/generate-cover-letter',
            'body': request.data,  # Your request payload here
            'headers': {
                'Content-Type': 'application/json'
            }
        }
    }

    # Calculate schedule time (optional)
    d = datetime.datetime.utcnow() + datetime.timedelta(seconds=10)
    timestamp = timestamp_pb2.Timestamp()
    timestamp.FromDatetime(d)
    task['schedule_time'] = timestamp

    # Add the task to the queue
    response = client.create_task(parent=parent, task=task)

    return jsonify({"status": "Task enqueued", "task_name": response.name})

@app.route('/generate-cover-letter', methods=['POST'])
def generate_cover_letter():
    try:
        print("Entering generate_cover_letter function...")
        data = get_data_from_request()
        validate_input_data(data)

        # Extract main user inputs
        email = data.get("email")
        name = data.get("name")
        agree_promo = data.get("agreePromo", False)

        user_doc_ref = db.collection('users').document(email)
        user_doc = user_doc_ref.get()
        num_letters = 0

        # If user exists, get the current num_letters
        if user_doc.exists:
            num_letters = user_doc.to_dict().get('num_letters', 0)

            # Check the limit
            if num_letters >= 5:
                raise Exception("Your free cover letters have run out")

        # Extract further user inputs for cover letter
        role = data.get("role")
        company = data.get("company")
        intro_points = data.get("intro_points")
        tailored_experience_points = data.get("tailored_experience_points")
        company_alignment_url = data.get("company_alignment_points")
        document_title = f"{role} at {company}"

        paragraphs = generate_paragraphs(name, role, company, intro_points, tailored_experience_points,
                                         company_alignment_url)
        output_stream = create_cover_letter(paragraphs, company, role)
        send_email_with_attachment(email, output_stream, document_title)

        # Update Firestore AFTER email sent successfully
        if user_doc.exists:
            update_data = {"num_letters": firestore.Increment(1)}
            current_agree_promo = user_doc.to_dict().get('agreePromo', False)
            if current_agree_promo != agree_promo:
                update_data['agreePromo'] = agree_promo
            user_doc_ref.update(update_data)
            remaining_articles = 5 - (num_letters + 1)
        else:
            user_doc_ref.set({
                "num_letters": 1,
                "agreePromo": agree_promo,
                "name": name
            })
            remaining_articles = 4

        return jsonify({
            "message": "Cover Letter successfully generated and sent!",
            "remaining_articles": remaining_articles
        }), 200

    except Exception as e:
        print(f"Error while generating cover letter: {str(e)}")
        return jsonify({"error": str(e)}), 500

def get_data_from_request():
    data = request.get_json()
    if not data:
        raise ValueError("No data received")
    return data

def validate_input_data(data):
    required_fields = ["name", "role", "company", "intro_points", "tailored_experience_points"]
    missing_fields = check_missing_fields(data, required_fields)
    if missing_fields:
        raise ValueError(f"Missing required fields: {', '.join(missing_fields)}")

def generate_paragraphs(name, role, company, intro_points, tailored_experience_points, company_alignment_url):
    intro_paragraph = rewrite_intro(name, role, company, intro_points)
    tailored_experience_paragraph = rewrite_tailored_experience(role, company, tailored_experience_points)
    company_alignment_paragraph = rewrite_company_alignment(company_alignment_url, name, company)
    return {
       "intro": intro_paragraph,
       "tailored_experience": tailored_experience_paragraph,
        "company_alignment": company_alignment_paragraph
    }

def create_cover_letter(paragraphs, company, role):
    intro_paragraph = paragraphs["intro"]
    tailored_experience_paragraph = paragraphs["tailored_experience"]
    company_alignment_paragraph = paragraphs["company_alignment"]

    document_title = f"Cover Letter for {role} at {company}"
    doc = Document()

    print("Initiated new Document")

    # Colored horizontal rule at the top
    section = doc.sections[0]
    section.top_margin = Pt(0.5)
    section.start_type
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # Add FULL NAME in 36pt font
    full_name_paragraph = doc.add_paragraph("[FULL NAME]")
    full_name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = full_name_paragraph.runs[0]
    run.font.size = Pt(36)

    # Add CITY, STATE, ZIP in normal font
    city_paragraph = doc.add_paragraph("[CITY STATE ZIP]")
    city_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add PHONE | EMAIL in normal font
    contact_paragraph = doc.add_paragraph("[PHONE] | [EMAIL]")
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # One newline
    # doc.add_paragraph("\n")

    # Add the paragraphs to the .docx document
    doc.add_paragraph(f"Dear Hiring Manager at {company},")
    # doc.add_paragraph("\n")
    doc.add_paragraph(intro_paragraph)
    # doc.add_paragraph("\n")
    doc.add_paragraph(tailored_experience_paragraph)
    # doc.add_paragraph("\n")
    doc.add_paragraph(company_alignment_paragraph)
    doc.add_paragraph(f"Sincerely,\n[FULL NAME]")

    # Save the .docx document
    file_name = f"{document_title}.docx"

    # Create the .docx document in memory using io.BytesIO
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)  # Reset the stream's position to the start

    return output_stream

# def store_user_data_in_firestore(name, email, agree_promo):
#
#     # Store the user's name and email in Firestore
#     users_ref = db.collection('users')  # 'users' is the collection name, you can change this if needed
#     _, doc_ref = users_ref.add({
#         'name': name,
#         'email': email,
#         'marketing_email_preference': agree_promo  # True if they agreed, False otherwise
#     })
#
#     if agree_promo:
#         print(f"Stored user data in Firestore with document ID: {doc_ref.id} and agreed to promotional emails.")
#     else:
#         print(f"Stored user data in Firestore with document ID: {doc_ref.id}. User did not agree to promotional emails.")


def send_email_with_attachment(email, output_stream, document_title):
    # Construct the message
    subject = "Your Generated Cover Letter"
    from_email = "noreply@mg.myaiguy.net"

    print(f"Preparing to send email to {email}")

    # Convert the BytesIO stream to bytes
    file_data = output_stream.getvalue()

    # Prepare the email data
    email_data = {
        "from": from_email,
        "to": email,
        "subject": subject,
        "text": "Please find your generated cover letter attached.",
        # You can enhance this with more meaningful text or even use HTML.
    }

    # Send the email using Mailgun's API
    response = requests.post(
        MAILGUN_API_ENDPOINT,
        auth=("api", MAILGUN_API_KEY),
        files=[("attachment", (f"{document_title}.docx", file_data,
                               'application/vnd.openxmlformats-officedocument.wordprocessingml.document'))],
        data=email_data
    )

    # Check for successful response
    if response.status_code == 200:
        print("Email sent successfully!")
    else:
        error_message = f"Error occurred while sending the email: {response.text}"
        print(error_message)
        raise Exception(error_message)

    # Error handling can be enhanced depending on requirements
    return response.json()

if __name__ == "__main__":
    app.run(debug=True)